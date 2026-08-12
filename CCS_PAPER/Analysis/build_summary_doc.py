"""
Builds a Word summary of the government-document text mining / sentiment / signal-term
analysis, for Britton to read later and pull into the CCS paper's background section.
"""
import csv
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

csv.field_size_limit(2**31 - 1)

ANALYSIS_DIR = r"C:\Users\britt\Desktop\Claude global\CCS PAPER\Analysis"
OUT_PATH = f"{ANALYSIS_DIR}\\Government_Document_Analysis_Summary.docx"

SIGNAL_CATEGORIES = [
    "Safety/risk assurance",
    "Risk/hazard language",
    "Transparency/accountability",
    "Community/consent signals",
    "Regulatory/authority signals",
    "Economic/benefit framing",
]


def read_csv(path):
    with open(path, encoding="utf-8") as f:
        return list(csv.DictReader(f))


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return table


def main():
    doc_sentiment = read_csv(f"{ANALYSIS_DIR}\\sentiment_by_document.csv")
    signal_terms = read_csv(f"{ANALYSIS_DIR}\\signal_terms_by_doc.csv")
    word_freq = read_csv(f"{ANALYSIS_DIR}\\word_frequency_overall.csv")

    document = Document()

    # Title
    title = document.add_heading("Louisiana CCS Government Document Corpus — Analysis Summary", level=0)

    meta = document.add_paragraph()
    meta.add_run("Corpus: Corpus_1 (18 documents, 414 pages, ~161,000 words)\n").italic = True
    meta.add_run("Prepared for: background/context material for the CCS paper — not Study 1 data.").italic = True

    document.add_paragraph()

    # Overview
    document.add_heading("1. What this is", level=1)
    document.add_paragraph(
        "This summarizes two automated passes over the Louisiana CCS legislative/regulatory "
        "corpus (bills, digests, fiscal notes, agency orders, a Federal Register notice, a public "
        "comment, and a Tunica Tribe document): (1) VADER sentiment scoring, and (2) a targeted "
        "keyword-frequency pass grouped into categories relevant to signaling theory and "
        "greenwashing/skepticism research (safety assurance, risk/hazard, transparency, "
        "community/consent, regulatory authority, and economic-benefit language)."
    )
    document.add_paragraph(
        "This corpus is government/regulatory text, not consumer or public discourse. It is not "
        "a substitute for Study 1 (which needs actual social media / public discourse data). Its "
        "value here is as institutional-signal context: characterizing what kind of signals the "
        "regulatory environment itself is sending, which motivates why consumer-facing skepticism "
        "and greenwashing perception are worth studying downstream."
    )

    # Key caveat callout
    p = document.add_paragraph()
    run = p.add_run(
        "Caveat: VADER is a lexicon-based sentiment tool built for informal/social text (tweets, "
        "reviews). It is a poor fit for formal legal/legislative prose — scores saturate near the "
        "extremes on long documents and should not be read as a validated measure of tone. The "
        "keyword-frequency pass is exploratory string-matching (no negation handling, unvalidated "
        "term list) — useful for framing and motivation, not for reporting as a validated coding "
        "scheme without a reliability check."
    )
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Section 2: Sentiment findings
    document.add_heading("2. Sentiment analysis (VADER, sentence-level)", level=1)
    document.add_paragraph(
        "Document-level compound scores saturate near +1.0 on anything longer than a page or two, "
        "so the more informative metric is the percentage of sentences scored positive/negative/"
        "neutral within each document. Sorted from most net-negative to most net-positive:"
    )

    rows = []
    sorted_rows = sorted(
        doc_sentiment,
        key=lambda r: float(r["pct_sentences_positive"] or 0) - float(r["pct_sentences_negative"] or 0),
    )
    for r in sorted_rows:
        rows.append([
            r["filename"],
            r["n_words"],
            f"{r['pct_sentences_positive']}%",
            f"{r['pct_sentences_negative']}%",
            f"{r['pct_sentences_neutral']}%",
        ])
    add_table(document, ["Document", "Words", "% Positive", "% Negative", "% Neutral"], rows)

    document.add_paragraph()
    document.add_heading("Notable documents", level=2)
    bullets = [
        ("Comment_1.pdf", "Most negative document (57% negative sentences). The one document in "
         "the corpus that is an outside/critical public voice rather than statutory or agency text."),
        ("HB_79_Damages_Threshold.pdf", "Second-most negative (40%), but see Section 3 — this reads "
         "as lexical saturation with risk/hazard/damage vocabulary (a bill defining damages "
         "thresholds), not necessarily negative emotional tone."),
        ("Tunica_Tribe.pdf", "Most positive document (86% positive sentences) and, per Section 3, by "
         "far the highest concentration of community/consent language in the corpus."),
        ("SB_244.pdf", "Largest document (227 pages, ~88,000 words); sentiment is diluted toward "
         "neutral/mixed as expected for a long enrolled bill covering many provisions."),
    ]
    for name, note in bullets:
        p = document.add_paragraph(style="List Bullet")
        p.add_run(f"{name}: ").bold = True
        p.add_run(note)

    # Section 3: Signal terms
    document.add_heading("3. Signaling-theory-relevant term frequency (per 1,000 words)", level=1)
    document.add_paragraph(
        "Categories are illustrative starting points, not a validated instrument. Values are raw "
        "keyword-match counts normalized per 1,000 words, so documents of different lengths are "
        "comparable."
    )

    rows = []
    for r in signal_terms:
        rows.append([r["filename"]] + [r[f"{cat} (per 1k words)"] for cat in SIGNAL_CATEGORIES])
    add_table(document, ["Document"] + [c.replace(" ", "\n") for c in SIGNAL_CATEGORIES], rows)

    document.add_paragraph()
    document.add_heading("Key patterns", level=2)
    patterns = [
        "Bills are heavy on regulatory/authority language (up to ~30 hits/1k words: commissioner, "
        "secretary, permit) but carry almost no safety-assurance or transparency language — "
        "statutory text defines who has authority; it doesn't make reassurance claims.",
        "Agency-level documents (Department Directive Order, Federal_Register.pdf) carry the "
        "corpus's highest concentrations of safety-assurance + transparency + regulatory-authority "
        "language together — the clearest examples of institutional reassurance signaling.",
        "Tunica_Tribe.pdf is a striking outlier: community/consent language at 38.9 per 1,000 "
        "words, more than 3x the next-highest document (HB_500_Mineral_Rights.pdf at 11.3) — the "
        "only document centered on consultation/consent rather than regulatory procedure.",
        "Economic/benefit framing is low corpus-wide except in fiscal-note-type documents "
        "(HB5_Fiscal.pdf, HB_79_Damages_Threshold.pdf), where it's definitionally expected.",
    ]
    for note in patterns:
        document.add_paragraph(note, style="List Bullet")

    # Section 4: Word frequency
    document.add_heading("4. Corpus-wide top terms", level=1)
    document.add_paragraph("Top 20 most frequent content words across the entire corpus (stopwords and legislative boilerplate removed):")
    top20 = word_freq[:20]
    rows = [[r["rank"], r["word"], r["count"]] for r in top20]
    add_table(document, ["Rank", "Term", "Count"], rows, col_widths=[0.8, 2.5, 1.2])

    # Section 5: Implications
    document.add_heading("5. Implications for the paper", level=1)
    document.add_paragraph(
        "This corpus is not Study 1 data — Study 1 needs actual public/consumer discourse "
        "(e.g., Reddit). This analysis is useful as institutional-signal context/motivation: it "
        "lets the paper state, with evidence, that the official regulatory discourse around CCS "
        "in Louisiana is dominated by regulatory-authority and procedural signaling, with safety-"
        "assurance and transparency language concentrated in agency orders rather than statute, "
        "and that community/consent language appears almost exclusively in the one Indigenous-"
        "authored document in the corpus. That asymmetry is a defensible motivation for why "
        "consumer-facing skepticism and greenwashing perception of CCS communication is worth "
        "studying downstream in Study 1/Study 2."
    )
    document.add_paragraph(
        "A data-quality note carried over from the original extraction: the Federal Register "
        "\"Class VI Primacy\" browser-printed PDF (doc 4) mostly failed to capture body text (a "
        "print-to-PDF artifact); the same content is covered cleanly by Federal_Register.pdf (doc 5)."
    )

    document.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    main()
