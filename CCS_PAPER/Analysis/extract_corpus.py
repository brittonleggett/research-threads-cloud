"""
Extract text from all PDFs (and docx) in Corpus_1 into a structured corpus CSV.
One row per (document, page) so page-level analysis is possible later;
a document-level rollup is written too.
"""
import csv
import re
from pathlib import Path

import pypdfium2 as pdfium
from docx import Document

CORPUS_DIR = Path(r"C:\Users\britt\Desktop\Claude global\CCS PAPER\Corpus_1")
OUT_DIR = Path(r"C:\Users\britt\Desktop\Claude global\CCS PAPER\Analysis")
OUT_DIR.mkdir(exist_ok=True)


def clean_text(text: str) -> str:
    text = text.replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_pdf(path: Path):
    rows = []
    pdf = pdfium.PdfDocument(str(path))
    try:
        for i, page in enumerate(pdf, start=1):
            textpage = page.get_textpage()
            text = textpage.get_text_range()
            textpage.close()
            page.close()
            rows.append((i, clean_text(text)))
    finally:
        pdf.close()
    return rows


def extract_docx(path: Path):
    doc = Document(str(path))
    text = "\n".join(p.text for p in doc.paragraphs)
    return [(1, clean_text(text))]


def main():
    doc_rows = []
    page_rows = []
    doc_id = 0

    files = sorted(CORPUS_DIR.iterdir())
    for path in files:
        if path.suffix.lower() == ".mqda":
            continue  # MAXQDA project file, not a source document
        if path.suffix.lower() not in (".pdf", ".docx"):
            continue

        doc_id += 1
        print(f"[{doc_id}] Extracting: {path.name}")
        try:
            if path.suffix.lower() == ".pdf":
                pages = extract_pdf(path)
            else:
                pages = extract_docx(path)
        except Exception as e:
            print(f"    ERROR extracting {path.name}: {e}")
            continue

        full_text_parts = []
        for page_num, text in pages:
            page_rows.append({
                "doc_id": doc_id,
                "filename": path.name,
                "page": page_num,
                "n_chars": len(text),
                "text": text,
            })
            full_text_parts.append(text)

        full_text = "\n\n".join(full_text_parts)
        doc_rows.append({
            "doc_id": doc_id,
            "filename": path.name,
            "n_pages": len(pages),
            "n_chars": len(full_text),
            "n_words": len(full_text.split()),
            "text": full_text,
        })

    # write page-level corpus
    page_csv = OUT_DIR / "corpus_pages.csv"
    with open(page_csv, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=["doc_id", "filename", "page", "n_chars", "text"])
        writer.writeheader()
        writer.writerows(page_rows)

    # write document-level corpus
    doc_csv = OUT_DIR / "corpus_documents.csv"
    with open(doc_csv, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=["doc_id", "filename", "n_pages", "n_chars", "n_words", "text"])
        writer.writeheader()
        writer.writerows(doc_rows)

    print(f"\nWrote {len(page_rows)} page rows -> {page_csv}")
    print(f"Wrote {len(doc_rows)} document rows -> {doc_csv}")


if __name__ == "__main__":
    main()
