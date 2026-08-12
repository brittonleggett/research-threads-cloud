import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

d = docx.Document()


def h1(text):
    d.add_heading(text, level=1)


def h2(text):
    d.add_heading(text, level=2)


def p(text, bold=False, italic=False):
    para = d.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    return para


title = d.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("MCNEESE STATE UNIVERSITY")
r.bold = True
r.font.size = Pt(16)

sub = d.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Human Subjects Institutional Review Board\nIRB APPLICATION PACKAGE (DRAFT v2)")
r.bold = True
r.font.size = Pt(13)

sub2 = d.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Consumer Responses to Tariff-Related Price Increase Messaging")
r.italic = True
r.font.size = Pt(12)

d.add_paragraph()
p("Principal Investigator: Britton R. Leggett")
p("Department: College of Business, McNeese State University")
p("Campus Address: Burton Business Center, 4205 Ryan Street, Lake Charles, LA 70605")
p("Email: bleggett1@mcneese.edu")
p("Phone: 337-475-5578 (ext. 5578)")
p("IRB Chair: Dr. Eddie Lyons | irbchair@mcneese.edu | 337-475-4077")
p("Date Prepared: 2026-08-04 (DRAFT -- pending PI review, CITI cert info, exact scale wording)")
p("Version: 2.0 -- supersedes the 4-cell single-factor design in v1.0")
d.add_page_break()

h1("Revision Note (v1.0 to v2.0)")
p("The prior draft used a single 4-level factor (silence / neutral tariff / gain-framed "
  "mitigation / loss-framed consequence, Study 2 target N=300). This version replaces it "
  "with a 3x2 between-subjects factorial -- Attribution Frame (tariff-explicit / "
  "general-cost-explicit / silent) x Cost-Response (full pass-through / shared-burden "
  "absorption) -- that Study 1's actual coded corpus themes independently validated. It "
  "also adds an explicit Pretest phase (N=150-180) that the prior draft did not cover, and "
  "updates the measured-construct list to a leaner model: Perceived Price Fairness, "
  "Perceived Opportunism, Trust in the Company, Purchase Intention, and Word-of-Mouth "
  "Intention. Fictional brand name (\"Meridian Home\") is unchanged for continuity.")

h1("Section 1: HSIRB Application Form Responses (Questions 1-10)")

h2("Question 1 -- Type of Project")
p("Academic faculty research project in the Department of Marketing, College of Business, "
  "McNeese State University. Not conducted in fulfillment of a course requirement; not "
  "sponsored by an external agency. Designed to generate findings for submission to the "
  "Journal of Consumer Marketing special issue \"Crafting Shape in a Fluid World\" "
  "(guest editors Kevin James, Janna Parker, Hyunju Shin), coordinated with the 2026 AMS "
  "Annual Conference.")

h2("Question 2 -- Funding")
p("Currently unfunded. No external grants, contracts, or sponsored research agreements. "
  "Participant compensation funded through the investigator's Prolific researcher account. "
  "If funding is obtained, an amendment will be submitted.")

h2("Question 3 -- Date to Begin Data Collection")
p("Anticipated start: as soon as written IRB exempt confirmation is received -- target "
  "2026-08-25 (placeholder; actual date depends on HSIRB turnaround, not yet confirmed). "
  "Study 1 does not involve human subjects and may proceed independently of this timeline; "
  "the investigator will still await written confirmation that Study 1 is excluded before "
  "treating it as formally cleared.")

h2("Question 4 -- Anticipated Date of Completion")
p("All three human-subjects phases (Pretest, then Study 2, then Study 3) targeted for "
  "completion by 2026-09-15. This is a tight sequence: the pretest must run, be analyzed, "
  "and (if needed) trigger vignette revisions before Study 2 can launch; only Study 3 can "
  "run in parallel with the other two, since it involves no experimental manipulation. If "
  "HSIRB turnaround runs past ~2026-08-25, this completion date compresses further.")

h2("Question 5 -- Status of Principal Investigator")
p("Faculty investigator: Britton R. Leggett, College of Business, Burton Business Center, "
  "4205 Ryan Street, Lake Charles, LA 70605. Phone: ext. 5578. Email: bleggett1@mcneese.edu. "
  "Faculty sponsor/supervisor: N/A (faculty-led project). CITI certification: [CERTIFICATE "
  "NUMBER / EXPIRATION -- pull from CITI records before submission].")

h2("Question 6 -- Exempt/Expedited Review and Rationale")
p("Study 1 does not constitute human subjects research under 45 CFR 46.102(e) and is "
  "excluded from IRB oversight (Chart 01, OHRP Common Rule Decision Charts). It analyzes "
  "only publicly available corporate communications (earnings calls, press coverage, price "
  "bulletins, SEC filings) made public by organizations for general audiences. No individual "
  "private data; no human contact. Secondary basis for exclusion if preferred: 45 CFR "
  "46.104(d)(4) (secondary research using publicly available information).")
p("Pretest, Study 2, and Study 3 all qualify for exemption under 45 CFR 46.104(d)(3): "
  "Benign Behavioral Interventions.", bold=True)
for item in [
    "Adult participants (18+) who prospectively agree to participate.",
    "Benign behavioral interventions: reading a brief, realistic business communication (or, "
    "for Study 3, completing a standard consumer-attitude survey with no experimental "
    "manipulation) and answering consumer-perception questions. Brief, harmless, not "
    "physically invasive, no reasonable expectation of significant adverse lasting impact.",
    "No identifying information collected; Prolific IDs used only for payment verification, "
    "stored separately, de-linked after payment.",
    "Mild incomplete disclosure (hypotheses not revealed in advance) remediated by a "
    "debriefing statement at the end of each survey.",
]:
    d.add_paragraph(item, style='List Bullet')

h2("Question 7 -- Abstract")
p("This research program examines how U.S. consumers respond to different corporate "
  "message strategies used to communicate tariff-driven price increases. As recently "
  "imposed import tariffs have led firms across consumer goods categories to announce price "
  "increases, a question of both academic and managerial importance has emerged: does the "
  "way a company frames a price-increase announcement shape consumer perceptions of "
  "fairness, trust, and behavioral intentions? Dual entitlement theory (Kahneman, Knetsch & "
  "Thaler, 1986) and the literature on inferred-motive price unfairness (Campbell, 1999) "
  "suggest that explicit, legitimate cost attribution is perceived as fairer than vague or "
  "absent explanation, while prospect theory (Kahneman & Tversky, 1979) and psychological "
  "reactance theory (Brehm, 1966) suggest that an unexplained, seemingly forced price change "
  "should provoke stronger negative reactions.")
p("The research program encompasses three sequential studies plus a stimulus pretest. Study "
  "1 (excluded from human-subjects review; public data only) employs thematic analysis of "
  "publicly available corporate communications about tariff-driven price increases, "
  "producing a coded typology of real-world messaging strategies that directly informs the "
  "manipulations tested below. A pretest (target N=150-180, Prolific) validates six "
  "candidate vignette stimuli for manipulation strength and absence of confounds (message "
  "length, perceived company likability, perceived realism) before they are used in the main "
  "study. Study 2 is a between-subjects online vignette experiment (target N=360-600) "
  "conducted via Prolific. Adult U.S. residents (18+) will be randomly assigned to one of six "
  "conditions in a 3 (Attribution Frame: tariff-explicit, general-cost-explicit, silent) x "
  "2 (Cost-Response: full pass-through, shared-burden absorption) factorial design. "
  "Participants will read a brief, realistic price-increase announcement from the fictional "
  "retailer \"Meridian Home\" and complete validated measures of perceived price "
  "fairness (Xia, Monroe & Cox, 2004), perceived opportunism/inferred motive (Campbell, "
  "1999), trust in the company (Chaudhuri & Holbrook, 2001), purchase intention (Dodds, "
  "Monroe & Grewal, 1991), and word-of-mouth intention (Maxham & Netemeyer, 2002). Study 3 "
  "is an independent online survey (target N=300-400) via Prolific to validate the full "
  "theoretical model via PLS-SEM, with no experimental manipulation. No vulnerable "
  "populations will be recruited. All data collection occurs online. Findings will be "
  "submitted to the Journal of Consumer Marketing special issue \"Crafting Shape in a "
  "Fluid World\" and presented at the 2026 AMS Annual Conference.")

h2("Question 8a -- Subject Population")
p("All \"No\" for every special/vulnerable population category (minors, pregnant "
  "women, prisoners, cognitively/decisionally impaired persons, economically disadvantaged "
  "persons specifically targeted, MSU students/students of the investigator, employees of "
  "the investigator, persons in dependent relationships with the investigator). Standard "
  "adult Prolific panel, screened for 18+, U.S. residence, English fluency.")

h2("Question 8b -- Number of Subjects")
for item in [
    "Study 1: 0 human subjects (public document analysis only).",
    "Pretest: up to 180.",
    "Study 2: up to 600.",
    "Study 3: up to 400.",
    "Total human subjects across the program: up to 1,180.",
]:
    d.add_paragraph(item, style='List Bullet')

h1("Recruitment, Compensation, Confidentiality, Consent, Risk/Benefit")

h2("Recruitment Procedure")
p("Participants for the Pretest, Study 2, and Study 3 will be recruited exclusively through "
  "Prolific (www.prolific.com). Eligibility: (1) age 18+, (2) U.S. residence, (3) English "
  "fluency. Prolific pre-screens on these attributes before showing the listing. Respondents "
  "are directed to a Qualtrics survey with informed consent and a brief eligibility "
  "screener up front; non-consenting or ineligible respondents are redirected out "
  "immediately.")

h2("Compensation")
p("Prolific academic-rate compensation (~$12.00/hour recommended pay), pro-rated by "
  "estimated completion time: Pretest (~6-7 min) ~$1.20-$1.40/participant; Study 2 (~12 min) "
  "~$2.00-$3.00/participant; Study 3 (~15 min) ~$2.50-$4.00/participant. Compensation is "
  "processed entirely through Prolific's infrastructure; the investigator does not have "
  "access to participant names or payment details.")

h2("Confidentiality")
p("No direct identifiers collected (no name, email, phone, IP address, or geolocation -- "
  "Qualtrics configured to suppress IP/location metadata). Prolific participant IDs used "
  "only to confirm completion and process payment, stored in a file separate from survey "
  "response data, permanently de-linked and deleted after payment confirmation. Data stored "
  "on password-protected, university-approved Qualtrics infrastructure during collection; "
  "exported and stored on a password-protected computer accessible only to the research team "
  "after collection. Retained a minimum of three years post-publication per APA guidelines, "
  "then securely deleted. All findings reported in aggregate only.")

h2("Voluntary Participation / Informed Consent")
p("Electronic informed consent at the start of each Qualtrics survey (~8th-grade reading "
  "level): study purpose, what participation involves, time commitment, risks/benefits, "
  "confidentiality, compensation, right to withdraw. \"I agree to participate\" / "
  "\"I do not agree to participate\" radio buttons; non-consenting respondents "
  "redirected out immediately, no penalty. Participants may withdraw at any point before "
  "final submission by closing the browser. Mild incomplete disclosure (hypotheses not "
  "revealed in advance) is remediated by a full debriefing statement at the end of each "
  "survey.")

h2("Risk/Benefit")
p("Minimal risk -- no greater than everyday life. Participants read a brief, realistic "
  "business announcement about a price increase (or, for Study 3, answer standard "
  "consumer-attitude questions) and respond to consumer-perception items. No physical, "
  "legal, financial, or lasting psychological risk; the only plausible risk is mild, "
  "transient annoyance at reading about a price increase. No sensitive personal information "
  "solicited beyond standard consumer-research demographics. Benefits are "
  "societal/scientific: empirical evidence on how corporate tariff-messaging strategy shapes "
  "fairness perceptions, trust, and purchase behavior. Risk/benefit ratio strongly "
  "favorable.")

h1("Section 6: Six Experimental Vignettes (Study 2)")
p("Fictional retailer: \"Meridian Home.\" Shared framing shown to all respondents "
  "before their assigned cell:", italic=True)
p("\"Imagine you are a regular customer of Meridian Home, a national retailer where you "
  "often shop. You receive the following email from the company: Subject: An Update on "
  "Pricing at Meridian Home\"")

vignettes = [
    ("Cell 1 -- Tariff-explicit x Full pass-through",
     "As you may have seen in the news, new tariffs on imported goods have significantly "
     "increased the cost of many of the products we sell. Because of these tariff-related "
     "costs, we are raising prices on select items starting next month. These increases "
     "reflect the full amount of the added tariff costs we are now paying to bring these "
     "products to you."),
    ("Cell 2 -- Tariff-explicit x Shared-burden",
     "As you may have seen in the news, new tariffs on imported goods have significantly "
     "increased the cost of many of the products we sell. Because of these tariff-related "
     "costs, we are raising prices on select items starting next month. We are absorbing a "
     "significant portion of these added tariff costs ourselves, and only passing along part "
     "of the increase to our customers."),
    ("Cell 3 -- General-cost-explicit x Full pass-through",
     "Due to rising costs across our supply chain, the cost of many of the products we sell "
     "has significantly increased. Because of these higher costs, we are raising prices on "
     "select items starting next month. These increases reflect the full amount of the added "
     "costs we are now paying to bring these products to you."),
    ("Cell 4 -- General-cost-explicit x Shared-burden",
     "Due to rising costs across our supply chain, the cost of many of the products we sell "
     "has significantly increased. Because of these higher costs, we are raising prices on "
     "select items starting next month. We are absorbing a significant portion of these "
     "added costs ourselves, and only passing along part of the increase to our customers."),
    ("Cell 5 -- Silent x Full pass-through",
     "We value you as a customer and want to keep you informed about changes at Meridian "
     "Home. Prices on select items will be increasing starting next month. This adjustment "
     "reflects the full increase in what it now costs us to bring these products to you. We "
     "appreciate your continued business and look forward to serving you."),
    ("Cell 6 -- Silent x Shared-burden",
     "We value you as a customer and want to keep you informed about changes at Meridian "
     "Home. Prices on select items will be increasing starting next month. We are absorbing "
     "a significant portion of the increase in what it now costs us to bring these products "
     "to you, and only passing along part of it. We appreciate your continued business and "
     "look forward to serving you."),
]
for title_txt, body_txt in vignettes:
    h2(title_txt)
    p(body_txt)

p("Open design question (flagged for PI sign-off, not yet resolved): the Silent condition "
  "(Cells 5-6) still references \"cost\" in a generic, unattributed way, since the "
  "Cost-Response factor needs something to manipulate in those two cells -- this is not "
  "true silence in the sense Study 1's real BMW artifact showed. See "
  "2026-08-04-vignette-drafts-v1.md for the two resolution options.", italic=True)

h1("Theoretical Model")
p("Diagram saved separately as Tariff_Model_Lean_2026-08-04.png (same folder). Path summary: "
  "Attribution Frame and Cost-Response (both manipulated) each feed Perceived Price Fairness "
  "and Perceived Opportunism; both mediators feed Trust in the Company; Trust feeds Purchase "
  "Intention and Word-of-Mouth Intention.")
try:
    d.add_picture("Tariff_Model_Lean_2026-08-04.png", width=docx.shared.Inches(6.3))
except Exception as e:
    p(f"[Diagram image not embedded -- insert Tariff_Model_Lean_2026-08-04.png manually. ({e})]")

h1("Scale Item Verification Status")
p("Searched for verbatim item wording for all five scales before drafting the instrument "
  "below. One came back fully verified; four did not and are flagged rather than presented "
  "as confirmed. Full detail in 2026-08-04-scale-items-verification-status.md.", italic=True)

h2("VERIFIED -- Trust in the Company (Chaudhuri & Holbrook, 2001)")
p("I trust this company. / I rely on this company. / This is an honest company. / This "
  "company is safe. (7-pt Likert, original alpha = .81)")

h2("UNVERIFIED -- Purchase Intention (Dodds, Monroe & Grewal, 1991) -- moderate-high confidence")
p("The likelihood of purchasing this product is high. / If I were going to buy a product "
  "like this, I would consider buying this one. / The probability that I would consider "
  "buying this product is high. (7-pt Likert) -- verify against original before fielding.")

h2("UNVERIFIED -- Perceived Price Fairness (Xia, Monroe & Cox, 2004) -- moderate confidence")
p("This price increase is: unfair-fair / unacceptable-acceptable / unreasonable-reasonable "
  "(7-pt semantic differential). Some published versions add a 4th item -- verify before "
  "fielding.")

h2("UNVERIFIED -- Word-of-Mouth Intention (Maxham & Netemeyer, 2002) -- moderate-low confidence")
p("I would say positive things about this company to other people. / I would recommend this "
  "company to someone who seeks my advice. / I would encourage friends and relatives to do "
  "business with this company. -- closely related to the Zeithaml, Berry & Parasuraman "
  "(1996) battery; verify which wording is actually Maxham & Netemeyer's before fielding.")

h2("UNVERIFIED -- Perceived Opportunism / Inferred Motive (Campbell, 1999) -- LOWEST confidence")
p("Illustrative only, not to be used as-is: \"I think the seller raised the price to take "
  "advantage of the situation\" / \"...to increase its profit margin\" / (reverse) "
  "\"...had a legitimate reason for raising the price.\" This is the most load-bearing scale "
  "in the design (tied directly to the Attribution Frame manipulation and Study 1's Theme 2) "
  "and the one with the least confident reconstruction -- verify this one first.")

h1("Full Survey Instrument -- Flow Order")

h2("Study 2 (vignette experiment)")
for item in [
    "Consent (~12 min estimate)",
    "Screener: age 18+, US residence, English fluency",
    "Random assignment to 1 of 6 vignette cells",
    "Vignette exposure (Section 6 above)",
    "Manipulation checks: attribution recall (multiple choice), attribution clarity "
    "(1-7), cost-response recall",
    "Perceived Price Fairness",
    "Perceived Opportunism",
    "Trust in the Company",
    "Purchase Intention",
    "Word-of-Mouth Intention",
    "Demographics: age, gender, income, education, shopping frequency, political "
    "ideology (7-pt)",
    "Debriefing statement (below)",
]:
    d.add_paragraph(item, style='List Number')

h2("Study 3 (SEM validation survey, no manipulation)")
for item in [
    "Consent (~15 min estimate)",
    "Screener (same as Study 2)",
    "Recall prompt: describe a real price increase from a retailer in the last 12 months "
    "(branch to generic-attitude version if nothing recalled)",
    "Same five scales, reworded to reference the recalled retailer/experience",
    "Demographics (same battery plus recency/category of recalled price increase)",
    "Shorter debriefing (no incomplete-disclosure element -- no manipulation to explain)",
]:
    d.add_paragraph(item, style='List Number')

h1("Debriefing Statement")
p("Thank you for completing this study.", bold=True)
p("We want to tell you a bit more about what this research is really about. This study "
  "examines how the way a company explains a price increase -- specifically, whether it "
  "names tariffs, cites general rising costs, or gives no explanation at all, and whether it "
  "says it's passing along the full cost or absorbing part of it -- affects how fair, "
  "trustworthy, and worth buying from consumers find that company. The message you read from "
  "\"Meridian Home\" was one of several different versions used in this study; Meridian Home "
  "is not a real company, and the email was written for research purposes only.")
p("We didn't tell you this in advance because knowing the specific comparison being tested "
  "can change how people respond, which would make the results less accurate. This is a "
  "standard and accepted research practice, not deception in the sense of anything false "
  "being told to you -- everything you were told about the study was true.")
p("Your responses are completely anonymous and cannot be linked back to you. If anything "
  "about this study concerns you, or if you'd like to withdraw your data even after "
  "completing the survey, you can contact the research team below.")
p("Researcher contact: Britton R. Leggett, bleggett1@mcneese.edu, 337-475-5578")
p("IRB contact: Dr. Eddie Lyons, IRB Chair, irbchair@mcneese.edu, 337-475-4077")

h1("Still Needed Before Submission-Ready")
for item in [
    "Verify the four [UNVERIFIED] scales against original source articles -- Campbell "
    "(1999) perceived-opportunism items first, they're the least certain and most "
    "load-bearing.",
    "CITI certificate number and expiration date.",
    "Build the actual Qualtrics file (randomization/branching/skip logic) from the flow "
    "order specified above.",
    "Sign-off on the placeholder 2026-08-25 start / 2026-09-15 completion dates once actual "
    "HSIRB turnaround time is known.",
    "PI decision on the Silent-condition design tension flagged in Section 6.",
]:
    d.add_paragraph(item, style='List Number')

d.save("IRB-Application-Package-Tariff-Messaging-v2-DRAFT.docx")
print("Saved.")
